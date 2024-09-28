import os
from utils.utils import load_json_file
from docx import Document

# Initialize a Word document
doc = Document()
doc.add_heading('Retrieval Task Results', level=1)

retrieval_tasks = [
    "minilm",
    "mpnet",
    "dpr",
    "contriever",
    "bge",
    "e5",
    "gte",
    "contriever_msmarco",
    "dragon_plus",
    "dragon_plus_self_rag",
]

dataset_paths = [
    "data/datasets/processed/medwiki.json",
    "data/datasets/processed/medquad.json",
]

for dataset_path in dataset_paths:
    # Add a heading for each dataset
    filename = dataset_path.split('/')[-1]  # Extract the filename, e.g., 'medwiki.json'
    dataset_name = filename.split('.')[0]  # Remove the extension, get 'medwiki'
    doc.add_heading(f'Results for {dataset_name}', level=2)

    # Create a table with 11 columns (Average rank, Top 0.2%, Top 0.5%, Top 0.8%, Top 1, Top 2, Top 3, Top 4, Top 5, Top 10, Top 20)
    # and len(retrieval_tasks) + 1 rows (one for headers, one for each retrieval task)
    table = doc.add_table(rows=1, cols=12)
    table.style = 'Table Grid'

    # Add the header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Model'
    hdr_cells[1].text = 'Average Rank'
    hdr_cells[2].text = 'Top 0.2%'
    hdr_cells[3].text = 'Top 0.5%'
    hdr_cells[4].text = 'Top 0.8%'
    hdr_cells[5].text = 'Top 1'
    hdr_cells[6].text = 'Top 2'
    hdr_cells[7].text = 'Top 3'
    hdr_cells[8].text = 'Top 4'
    hdr_cells[9].text = 'Top 5'
    hdr_cells[10].text = 'Top 10'
    hdr_cells[11].text = 'Top 20'

    for i, retrieval_task in enumerate(retrieval_tasks):
        result_path = f"data/results/{dataset_name}_{retrieval_task}.json"
        result = load_json_file(result_path)
        len_result = len(result)

        rank_sum = 0
        top_percent = [0] * 10  # Store the top percentages

        for item in result:
            rank_sum += item["rank"] + 1
            if item["rank"] < len_result * 0.002:
                top_percent[0] += 1
            if item["rank"] < len_result * 0.005:
                top_percent[1] += 1
            if item["rank"] < len_result * 0.008:
                top_percent[2] += 1
            if item["rank"] < 1:
                top_percent[3] += 1
            if item["rank"] < 2:
                top_percent[4] += 1
            if item["rank"] < 3:
                top_percent[5] += 1
            if item["rank"] < 4:
                top_percent[6] += 1
            if item["rank"] < 5:
                top_percent[7] += 1
            if item["rank"] < 10:
                top_percent[8] += 1
            if item["rank"] < 20:
                top_percent[9] += 1

        # Calculate the metrics
        avg_rank = rank_sum / len_result
        top_0_2 = top_percent[0] / len_result * 100
        top_0_5 = top_percent[1] / len_result * 100
        top_0_8 = top_percent[2] / len_result * 100
        top_1 = top_percent[3] / len_result * 100
        top_2 = top_percent[4] / len_result * 100
        top_3 = top_percent[5] / len_result * 100
        top_4 = top_percent[6] / len_result * 100
        top_5 = top_percent[7] / len_result * 100
        top_10 = top_percent[8] / len_result * 100
        top_20 = top_percent[9] / len_result * 100

        # Fill the table row
        row_cells = table.add_row().cells
        row_cells[0].text = retrieval_task
        row_cells[1].text = f"{avg_rank:.2f}"
        row_cells[2].text = f"{top_0_2:.2f}%"
        row_cells[3].text = f"{top_0_5:.2f}%"
        row_cells[4].text = f"{top_0_8:.2f}%"
        row_cells[5].text = f"{top_1:.2f}%"
        row_cells[6].text = f"{top_2:.2f}%"
        row_cells[7].text = f"{top_3:.2f}%"
        row_cells[8].text = f"{top_4:.2f}%"
        row_cells[9].text = f"{top_5:.2f}%"
        row_cells[10].text = f"{top_10:.2f}%"
        row_cells[11].text = f"{top_20:.2f}%"

    doc.add_paragraph()  # Add a paragraph to separate tables

# Save the document
doc.save('retrieval_task_results.docx')
